from docx import Document
from pathlib import Path

content = {
    "title": "TECHPULSE-AI",
    "subtitle": "Assistant intelligent de cybersécurité pour les agences de voyage",
    "sections": [
        (
            "1. Problématique",
            "Les agences de voyage utilisent quotidiennement des systèmes de réservation, des plateformes de billetterie, des solutions de paiement en ligne et des outils de gestion des clients. Ces technologies sont exposées à des vulnérabilités de sécurité, des cyberattaques et des tentatives de fraude.\n\nLa majorité des petites et moyennes agences ne disposent pas d'une équipe spécialisée en cybersécurité capable de surveiller en permanence les nouvelles menaces. Une faille non détectée ou une attaque réussie peut entraîner des pertes financières, une interruption de service et une atteinte à la réputation de l'entreprise.",
        ),
        (
            "2. Solution Proposée",
            "TECHPULSE-AI est un assistant intelligent de cybersécurité conçu spécialement pour les agences de voyage.\n\nLa plateforme collecte et analyse automatiquement les informations de sécurité provenant de sources fiables afin d'identifier les menaces pouvant affecter les systèmes de réservation et de billetterie.\n\nGrâce à l'intelligence artificielle, elle fournit des alertes, des recommandations de correction et un assistant conversationnel permettant aux utilisateurs de comprendre rapidement les risques et les actions à entreprendre.",
        ),
        (
            "3. Objectifs",
            "- Surveiller les vulnérabilités affectant les technologies utilisées dans le secteur du voyage.\n- Fournir des alertes de sécurité pertinentes.\n- Aider les agences à mieux comprendre les risques auxquels elles sont exposées.\n- Générer des recommandations de correction adaptées.\n- Simplifier l'accès aux informations de cybersécurité grâce à l'intelligence artificielle.",
        ),
        (
            "4. Utilisateurs Cibles",
            "- Agences de voyage\n- Agences de billetterie\n- Responsables informatiques\n- Administrateurs systèmes\n- Responsables de la sécurité informatique\n- Gestionnaires de plateformes de réservation",
        ),
        (
            "5. Fonctionnalités MVP",
            "- Surveillance des vulnérabilités de sécurité.\n- Classification automatique des risques.\n- Recherche intelligente dans les données de cybersécurité.\n- Assistant IA conversationnel.\n- Tableau de bord de suivi des menaces.\n- Recommandations de correction.",
        ),
        (
            "6. Fonctionnalités Futures",
            "- Alertes de sécurité en temps réel.\n- Analyse des incidents récents du secteur du voyage.\n- Analyse des risques de fraude liés à la billetterie.\n- Historique des vulnérabilités surveillées.\n- Génération automatique de rapports de sécurité.\n- Recommandations de prévention contre la fraude.\n- Détection intelligente des comportements suspects.",
        ),
        (
            "7. Architecture Technique",
            "1. Collecte des données via des API de cybersécurité.\n2. Nettoyage et prétraitement des données.\n3. Classification des vulnérabilités à l'aide de modèles NLP.\n4. Génération d'embeddings vectoriels.\n5. Stockage dans une base vectorielle FAISS.\n6. Recherche sémantique et système RAG.\n7. Génération de réponses via un assistant IA.\n8. Visualisation dans un tableau de bord interactif.",
        ),
        (
            "8. Technologies Utilisées",
            "- Python\n- APIs\n- JSON\n- NLP (Traitement Automatique du Langage Naturel)\n- Vector Embeddings\n- FAISS\n- Retrieval-Augmented Generation (RAG)\n- Transformers\n- Machine Learning Supervisé\n- Feature Engineering",
        ),
        (
            "9. Valeur Ajoutée",
            "TECHPULSE-AI permet aux agences de voyage d'accéder à une expertise cybersécurité sans disposer d'une équipe spécialisée. La solution centralise les informations critiques, simplifie l'analyse des menaces et aide les entreprises à prendre rapidement les bonnes décisions pour protéger leurs activités et leurs clients.",
        ),
        (
            "10. Vision à Long Terme",
            "Faire de TECHPULSE-AI une plateforme de référence pour la cybersécurité des agences de voyage, capable de surveiller les menaces, anticiper les risques de fraude, assister les équipes métier et renforcer la sécurité des systèmes de réservation et de billetterie.",
        ),
    ],
}

doc = Document()
doc.add_heading(content['title'], level=0)
paragraph = doc.add_paragraph(content['subtitle'])
paragraph.runs[0].italic = True
for title, body in content['sections']:
    doc.add_heading(title, level=1)
    chunks = body.split('\n\n')
    for chunk in chunks:
        if chunk.startswith('- '):
            for line in chunk.split('\n'):
                doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(chunk)

output_path = Path('docs/TECHPULSE-AI.docx')
output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(output_path)
print(f'Wrote {output_path}')
